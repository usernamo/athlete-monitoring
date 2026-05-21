# -*- coding: utf-8 -*-
"""Analyze athlete diary Excel/DOCX sources. Run: python scripts/analyze_sources.py"""
import os
from io import StringIO

import pandas as pd
from openpyxl import load_workbook
from docx import Document

BASE = r"c:\Users\PC\Downloads\Telegram Desktop"
OUT = os.path.join(os.path.dirname(__file__), "..", "source-analysis.txt")

XLSX_FILES = [
    os.path.join(BASE, "дневник спортсмена 1.xlsx"),
    os.path.join(BASE, "дневник спортсмена 2.xlsx"),
    os.path.join(BASE, "дневник спортсмена.xlsx"),
]
DOCX_FILE = os.path.join(BASE, "Дневник спортсмена (общ).docx")


def df_to_text_table(df, max_rows=20):
    view = df.head(max_rows)
    buf = StringIO()
    view.to_string(buf=buf, index=True, max_cols=50, max_colwidth=40)
    return buf.getvalue()


def analyze_xlsx(path, lines):
    lines.append("=" * 80)
    lines.append(f"FILE: {path}")
    lines.append("=" * 80)
    if not os.path.isfile(path):
        lines.append("ERROR: File not found")
        lines.append("")
        return
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        lines.append(f"Sheet count: {len(wb.sheetnames)}")
        lines.append(f"Sheet names: {wb.sheetnames}")
        lines.append("")
        wb.close()

        xl = pd.ExcelFile(path, engine="openpyxl")
        for name in xl.sheet_names:
            lines.append("-" * 60)
            lines.append(f"SHEET: {name!r}")
            df = pd.read_excel(path, sheet_name=name, header=None, engine="openpyxl")
            nrows, ncols = df.shape
            lines.append(f"Dimensions: {nrows} rows x {ncols} columns")
            lines.append("First 20 rows (text table):")
            lines.append(df_to_text_table(df, 20))
            lines.append("")
        xl.close()
    except Exception as e:
        lines.append(f"ERROR reading xlsx: {type(e).__name__}: {e}")
        lines.append("")


def analyze_docx(path, lines):
    lines.append("=" * 80)
    lines.append(f"FILE: {path}")
    lines.append("=" * 80)
    if not os.path.isfile(path):
        lines.append("ERROR: File not found")
        lines.append("")
        return
    try:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        full = "\n".join(parts)
        lines.append(f"Paragraph count: {len(doc.paragraphs)}")
        lines.append(f"Table count: {len(doc.tables)}")
        lines.append(f"Extracted text length: {len(full)} characters")
        lines.append("")
        lines.append("--- FULL TEXT ---")
        lines.append(full)
        lines.append("")
    except Exception as e:
        lines.append(f"ERROR reading docx: {type(e).__name__}: {e}")
        lines.append("")


def main():
    lines = ["SOURCE ANALYSIS: Athlete diary files", f"Base: {BASE}", ""]
    for p in XLSX_FILES:
        analyze_xlsx(p, lines)
    analyze_docx(DOCX_FILE, lines)
    text = "\n".join(lines)
    out_path = os.path.normpath(OUT)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"Wrote {out_path} ({len(text)} chars)")


if __name__ == "__main__":
    main()
